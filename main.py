import stanford_ai
import arxiv_scraper
import techcrunch
import venture_beat
import theverge
import thn
from docx import Document
from datetime import datetime
import os
import logging
import db_manager

OUTPUT_DIR = 'summary'

def setup_environment():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    db_manager.init_db()

def export_all_summaries(arxiv_data=None, techcrunch_data=None, venturebeat_data=None, stanford_data=None, theverge_data=None, thn_data=None):
    """Exportiert alle Zusammenfassungen in eine einzige Word-Datei"""
    try:
        doc = Document()
        today_date = datetime.today().strftime('%d. %B %Y')
        doc.add_heading(f'AI Research Zusammenfassungen - {today_date}', level=0)
        
        # ArXiv Abschnitt
        if arxiv_data:
            doc.add_heading('ArXiv Forschungsarbeiten', level=1)
            for title, data in arxiv_data.items():
                doc.add_heading(title, level=2)
                doc.add_paragraph(f"Autoren: {data['authors']}")
                doc.add_paragraph(f"Veröffentlicht: {data['published']}")
                doc.add_paragraph(f"Link: {data['link']}", style='Intense Quote')

                # Add citations if available
                # if 'citations' in data and data['citations']:
                #     doc.add_heading('Zitationen:', level=3)
                #     for style, citation_text in data['citations'].items():
                #         doc.add_paragraph(f"{style.upper()}: {citation_text}")
                        
                doc.add_paragraph(data['summary'])
                doc.add_paragraph("-" * 80)
        
        # TechCrunch Abschnitt
        if techcrunch_data:
            doc.add_heading('TechCrunch Artikel', level=1)
            for article in techcrunch_data:
                doc.add_heading(article["Titel"], level=2)
                doc.add_paragraph(f"Datum: {article['Datum']}")
                doc.add_paragraph(f"Link: {article['Link']}", style='Intense Quote')
                doc.add_paragraph(article['Zusammenfassung'])
                doc.add_paragraph("-" * 80)
        
        # VentureBeat Abschnitt
        if venturebeat_data:
            doc.add_heading('VentureBeat Artikel', level=1)
            for article in venturebeat_data:
                doc.add_heading(article["title"], level=2)
                # doc.add_paragraph(f"Autor: {article['author']}")
                doc.add_paragraph(f"Datum: {article['pub_date'].strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph(f"Link: {article['link']}", style='Intense Quote')
                doc.add_paragraph(article["summary"])
                doc.add_paragraph("-" * 80)
        
        # Stanford Abschnitt
        if stanford_data:
            doc.add_heading('Stanford Blog Artikel', level=1)
            for title, link, summary in stanford_data:
                doc.add_heading(title, level=2)
                doc.add_paragraph(f"Link: {link}", style='Intense Quote')
                doc.add_paragraph(summary)
                doc.add_paragraph("-" * 80)

        # TheVerge Abschnitt
        if theverge_data:
            doc.add_heading('The Verge Artikel', level=1)
            for article in theverge_data:
                doc.add_heading(article["Titel"], level=2)
                doc.add_paragraph(f"Datum: {article['Datum']}")
                doc.add_paragraph(f"Link: {article['Link']}", style='Intense Quote')
                doc.add_paragraph(article['Zusammenfassung'])
                doc.add_paragraph("-" * 80)
        
        # TheHackerNews Abschnitt
        if thn_data:
            doc.add_heading('The Hacker News Artikel', level=1)
            for article in thn_data:
                doc.add_heading(article["Titel"], level=2)
                doc.add_paragraph(f"Datum: {article['Datum']}")
                doc.add_paragraph(f"Link: {article['Link']}", style='Intense Quote')
                doc.add_paragraph(article['Zusammenfassung'])
                doc.add_paragraph("-" * 80)
        
        # Dokument speichern
        filename = f"AI_Zusammenfassungen_{datetime.today().strftime('%Y-%m-%d')}.docx"
        doc.save(os.path.join(OUTPUT_DIR, filename))
        logging.info(f"Alle Zusammenfassungen wurden in {filename} gespeichert")
        
    except Exception as e:
        logging.error(f"Fehler beim Export der Zusammenfassungen: {str(e)}")

if __name__ == "__main__":
    setup_environment()
    
    # Get citation styles (can be modified to accept command line arguments)
    citation_styles = ["apa", "mla", "chicago", "bibtex"]
    
    # Alle Daten sammeln
    arxiv_data = arxiv_scraper.main_arxiv(citation_styles=citation_styles)
    techcrunch_data = techcrunch.main_techcrunch()
    venturebeat_data = venture_beat.main_venturebeat()
    stanford_data = stanford_ai.main_stanford()
    theverge_data = theverge.main_verge()
    thn_data = thn.main_thn()
    
    # Alle Daten in einer Datei speichern
    export_all_summaries(
        arxiv_data=arxiv_data,
        techcrunch_data=techcrunch_data,
        venturebeat_data=venturebeat_data,
        stanford_data=stanford_data,
        theverge_data=theverge_data,
        thn_data=thn_data
    )
    
    print("✅ Programm erfolgreich abgeschlossen")
